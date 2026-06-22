"""
Profile Parser Agent — Stage 2
Uses LlamaIndex + Groq to extract structured JSON from every resume.
Falls back to PyMuPDF + regex extraction when LlamaIndex is unavailable.
Fires profile.parsed event on completion.
"""
import json
import re
import os
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from shared.db.database import get_db, generate_id
from shared.db.models import Candidate, Application, AuditLog
from shared.queue.event_bus import event_bus
from shared.queue.event_topics import EventTopics
from shared.auth.jwt_middleware import get_current_user
from config import GROQ_API_KEY, LLAMA_INDEX_CHUNK_SIZE, LLAMA_INDEX_CHUNK_OVERLAP

router = APIRouter(prefix="/sourcing", tags=["Candidate Intake — Stage 2"])


class ParseRequest(BaseModel):
    candidate_id: str


class ParseBatchRequest(BaseModel):
    candidate_ids: list[str] = []
    parse_all_unparsed: bool = False


# ── Text Extraction ───────────────────────────────────

def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"⚠️ PDF extraction error: {e}")
        return ""


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text.strip()
    except Exception as e:
        print(f"⚠️ DOCX extraction error: {e}")
        return ""


def _extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        print(f"⚠️ TXT extraction error: {e}")
        return ""


def extract_text(file_path: str) -> str:
    """Extract text from any supported file type."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_text_from_docx(file_path)
    elif ext == ".txt":
        return _extract_text_from_txt(file_path)
    else:
        return ""


# ── LlamaIndex + Groq Parsing ───────────────────────

async def _parse_with_llamaindex(text: str) -> dict:
    """
    Use LlamaIndex with Groq to extract structured resume data.
    This is the primary parsing method.
    """
    if not GROQ_API_KEY:
        print("⚠️ No Groq API key, falling back to regex parsing")
        return _parse_with_regex(text)

    try:
        from llama_index.core import Document, VectorStoreIndex, Settings
        from llama_index.core.node_parser import SentenceSplitter
        from llama_index.llms.groq import Groq

        print("🧠 Using LlamaIndex + Groq for resume parsing...")

        # Configure LlamaIndex with Groq (using Llama 3.1 8B)
        llm = Groq(
            model="llama-3.1-8b-instant",
            api_key=GROQ_API_KEY,
            temperature=0.1
        )
        Settings.llm = llm
        Settings.chunk_size = LLAMA_INDEX_CHUNK_SIZE
        Settings.chunk_overlap = LLAMA_INDEX_CHUNK_OVERLAP

        # Create document and index
        documents = [Document(text=text)]
        node_parser = SentenceSplitter(
            chunk_size=LLAMA_INDEX_CHUNK_SIZE, 
            chunk_overlap=LLAMA_INDEX_CHUNK_OVERLAP
        )
        nodes = node_parser.get_nodes_from_documents(documents)
        
        # Build vector index
        index = VectorStoreIndex(nodes)
        
        # Use simple query engine
        query_engine = index.as_query_engine(
            response_mode="compact",
            similarity_top_k=2
        )

        # Optimized extraction prompt for Groq/Llama
        extraction_prompt = """Extract resume information as JSON:

{
    "name": "candidate full name",
    "email": "email address or null",
    "phone": "phone number or null", 
    "location": "city, country or null",
    "current_role": "current job title or null",
    "experience_years": 0,
    "summary": "brief professional summary",
    "skills": ["skill1", "skill2"],
    "projects": [{"name": "project name", "description": "details and tech stack/tools used"}],
    "education": [{"degree": "degree", "institution": "school", "year": "year"}],
    "work_history": [{"company": "company", "role": "title", "duration": "period"}],
    "certifications": ["cert1"],
    "languages": ["lang1"]
}

Return only valid JSON, no other text."""

        print("🔍 Querying LlamaIndex + Groq for structured data extraction...")
        response = query_engine.query(extraction_prompt)
        response_text = str(response).strip()

        print(f"📄 Groq response: {response_text[:200]}...")

        # Try to extract JSON from response
        try:
            # Look for JSON object in response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_str = response_text[json_start:json_end]
                parsed = json.loads(json_str)
                
                # Ensure required fields exist
                parsed.setdefault("name", "Unknown")
                parsed.setdefault("skills", [])
                parsed.setdefault("projects", [])
                parsed.setdefault("education", [])
                parsed.setdefault("work_history", [])
                parsed.setdefault("certifications", [])
                parsed.setdefault("languages", [])
                
                parsed["_parser"] = "llamaindex_groq"
                print(f"✅ LlamaIndex + Groq parsing successful - extracted {len(parsed.get('skills', []))} skills")
                return parsed
            else:
                raise ValueError("No JSON object found in response")
                
        except (json.JSONDecodeError, ValueError) as e:
            print(f"⚠️ JSON parsing failed: {e}")
            print(f"Raw response: {response_text}")

        # If LlamaIndex extraction failed, try direct Groq call
        print("⚠️ LlamaIndex JSON extraction failed, falling back to direct Groq")
        return await _parse_with_groq_direct(text)

    except ImportError as e:
        print(f"⚠️ LlamaIndex Groq import error: {e}")
        return await _parse_with_groq_direct(text)
    except Exception as e:
        print(f"⚠️ LlamaIndex + Groq parsing error: {e}")
        import traceback
        traceback.print_exc()
        return await _parse_with_groq_direct(text)


async def _parse_with_groq_direct(text: str) -> dict:
    """
    Direct Groq API call for resume parsing (fallback method).
    """
    if not GROQ_API_KEY:
        print("⚠️ No Groq API key, falling back to regex parsing")
        return _parse_with_regex(text)
    
    try:
        from groq import AsyncGroq
        
        print("🔄 Using direct Groq API for resume parsing...")
        
        client = AsyncGroq(api_key=GROQ_API_KEY)
        
        prompt = f"""Extract structured information from this resume text and return ONLY a valid JSON object:

Resume Text:
{text[:4000]}  

Return JSON with these exact keys:
{{
    "name": "full name or null",
    "email": "email or null",
    "phone": "phone or null",
    "location": "location or null", 
    "current_role": "job title or null",
    "experience_years": 0,
    "summary": "brief summary",
    "skills": ["skill1", "skill2"],
    "projects": [{{"name": "project name", "description": "details and tech stack/tools used"}}],
    "education": [{{"degree": "degree", "institution": "school", "year": "year"}}],
    "work_history": [{{"company": "company", "role": "title", "duration": "period"}}],
    "certifications": ["cert1"],
    "languages": ["lang1"]
}}

Return ONLY the JSON object, no other text."""

        response = await client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": "You are an expert resume parser. Extract information and return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content.strip()
        print(f"📄 Direct Groq response: {response_text[:200]}...")
        
        # Extract JSON from response
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        
        if json_start >= 0 and json_end > json_start:
            json_str = response_text[json_start:json_end]
            parsed = json.loads(json_str)
            
            # Ensure required fields
            parsed.setdefault("name", "Unknown")
            parsed.setdefault("skills", [])
            parsed.setdefault("projects", [])
            parsed.setdefault("education", [])
            parsed.setdefault("work_history", [])
            parsed.setdefault("certifications", [])
            parsed.setdefault("languages", [])
            
            parsed["_parser"] = "groq_direct"
            print(f"✅ Direct Groq parsing successful")
            return parsed
            
    except Exception as e:
        print(f"⚠️ Direct Groq error: {e}")
    
    # Final fallback to regex
    print("⚠️ All AI parsing failed, falling back to regex")
    return _parse_with_regex(text)


async def _parse_with_ollama_direct(text: str) -> dict:
    """
    Direct Ollama API call for resume parsing (fallback method).
    """
    try:
        import httpx
        
        print("🔄 Using direct Ollama API for resume parsing...")
        
        prompt = f"""Extract structured information from this resume text and return ONLY a valid JSON object:

Resume Text:
{text[:3000]}  

Return JSON with these exact keys:
{{
    "name": "full name or null",
    "email": "email or null",
    "phone": "phone or null",
    "location": "location or null", 
    "current_role": "job title or null",
    "experience_years": 0,
    "summary": "brief summary",
    "skills": ["skill1", "skill2"],
    "projects": [{{"name": "project name", "description": "details and tech stack/tools used"}}],
    "education": [{{"degree": "degree", "institution": "school", "year": "year"}}],
    "work_history": [{{"company": "company", "role": "title", "duration": "period"}}],
    "certifications": ["cert1"],
    "languages": ["lang1"]
}}

Return ONLY the JSON object, no other text."""

        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": "llama3.2:3b",
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,
                        "top_p": 0.9
                    }
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                response_text = result.get("response", "").strip()
                
                print(f"📄 Direct Ollama response: {response_text[:200]}...")
                
                # Extract JSON from response
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                
                if json_start >= 0 and json_end > json_start:
                    json_str = response_text[json_start:json_end]
                    parsed = json.loads(json_str)
                    
                    # Ensure required fields
                    parsed.setdefault("name", "Unknown")
                    parsed.setdefault("skills", [])
                    parsed.setdefault("projects", [])
                    parsed.setdefault("education", [])
                    parsed.setdefault("work_history", [])
                    parsed.setdefault("certifications", [])
                    parsed.setdefault("languages", [])
                    
                    parsed["_parser"] = "ollama_direct"
                    print(f"✅ Direct Ollama parsing successful")
                    return parsed
                    
            print(f"⚠️ Ollama API error: {response.status_code}")
            
    except Exception as e:
        print(f"⚠️ Direct Ollama error: {e}")
    
    # Final fallback to regex
    print("⚠️ All AI parsing failed, falling back to regex")
    return _parse_with_regex(text)


# ── Regex Fallback Parser ─────────────────────────────

def _parse_with_regex(text: str) -> dict:
    """Fallback parser using regex patterns when AI is unavailable."""
    # Extract email
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.]+', text)
    email = email_match.group() if email_match else None

    # Extract phone
    phone_match = re.search(r'[\+]?[\d\s\-\(\)]{10,15}', text)
    phone = phone_match.group().strip() if phone_match else None

    # Extract name (usually the first line)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    name = lines[0] if lines else "Unknown"
    # If first line looks like an email or phone, try next line
    if name and ('@' in name or re.match(r'^[\d\+\-\(\)\s]+$', name)):
        name = lines[1] if len(lines) > 1 else "Unknown"

    # Extract skills (common patterns)
    skill_keywords = [
        "Python", "Java", "JavaScript", "TypeScript", "React", "Angular", "Vue",
        "Node.js", "Django", "Flask", "FastAPI", "Spring", "Docker", "Kubernetes",
        "AWS", "Azure", "GCP", "SQL", "PostgreSQL", "MongoDB", "Redis",
        "TensorFlow", "PyTorch", "Machine Learning", "Deep Learning",
        "Git", "CI/CD", "Linux", "Agile", "Scrum", "REST API",
        "HTML", "CSS", "C++", "C#", "Go", "Rust", "Ruby", "PHP",
    ]
    found_skills = [s for s in skill_keywords if s.lower() in text.lower()]

    # Extract education
    education = []
    edu_patterns = re.findall(
        r'(B\.?Tech|B\.?E|M\.?Tech|M\.?S|MBA|Ph\.?D|B\.?Sc|M\.?Sc|Bachelor|Master)[\s\w,\-]*(?:from|at|,)\s*([\w\s]+)',
        text, re.IGNORECASE
    )
    for degree, inst in edu_patterns:
        education.append({"degree": degree.strip(), "institution": inst.strip(), "year": ""})

    # Extract experience years
    exp_match = re.search(r'(\d+)\+?\s*(?:years?|yrs?)[\s]*(?:of)?\s*(?:experience)?', text, re.IGNORECASE)
    experience_years = int(exp_match.group(1)) if exp_match else 0

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "location": None,
        "current_role": None,
        "experience_years": experience_years,
        "summary": text[:300] + "..." if len(text) > 300 else text,
        "skills": found_skills,
        "projects": [],
        "education": education,
        "work_history": [],
        "certifications": [],
        "languages": [],
        "_parser": "regex_fallback",
    }


# ── Endpoints ──────────────────────────────────────────

async def _parse_uploaded_resume_internal(candidate_id: str):
    """Internal function to parse a resume automatically when uploaded."""
    from shared.db.database import async_session
    from shared.db.models import Candidate, Application, AuditLog
    from sqlalchemy import select
    import os
    import json
    from datetime import datetime
    
    async with async_session() as db:
        try:
            # Get candidate and associated application
            result = await db.execute(select(Candidate).where(Candidate.id == candidate_id))
            candidate = result.scalar_one_or_none()
            if not candidate:
                print(f"❌ Candidate {candidate_id} not found for auto-parsing")
                return

            # Get application if exists
            app_result = await db.execute(select(Application).where(Application.candidate_id == candidate_id))
            application = app_result.scalar_one_or_none()
            job_id = application.job_id if application else candidate.job_id

            if not candidate.resume_url or not os.path.exists(candidate.resume_url):
                print(f"❌ No resume file found for candidate {candidate_id}")
                return

            # Step 1: Extract text
            raw_text = extract_text(candidate.resume_url)
            if not raw_text:
                print(f"❌ Could not extract text from resume for candidate {candidate_id}")
                return

            # Step 2: Parse with LlamaIndex
            parsed_data = await _parse_with_llamaindex(raw_text)

            # Step 3: Update candidate record
            candidate.raw_resume_text = raw_text
            candidate.parsed_data = json.dumps(parsed_data)
            candidate.name = parsed_data.get("name") or candidate.name
            candidate.email = parsed_data.get("email") or candidate.email
            candidate.phone = parsed_data.get("phone") or candidate.phone
            candidate.location = parsed_data.get("location") or candidate.location
            candidate.current_role = parsed_data.get("current_role") or candidate.current_role
            candidate.experience_years = parsed_data.get("experience_years") or candidate.experience_years
            candidate.skills = json.dumps(parsed_data.get("skills", []))
            candidate.education = json.dumps(parsed_data.get("education", []))
            candidate.work_history = json.dumps(parsed_data.get("work_history", []))
            candidate.status = "parsed"
            candidate.updated_at = datetime.utcnow()
            
            # CRITICAL FIX: Set job_id on candidate for Stage 3 screening
            if job_id:
                candidate.job_id = job_id
                print(f"✅ Set job_id {job_id} on candidate {candidate_id}")
            else:
                print(f"⚠️ No job_id found for candidate {candidate_id} - Stage 3 screening will fail")

            # Audit log
            audit = AuditLog(
                id=generate_id(),
                event_type=EventTopics.PROFILE_PARSED,
                agent_name="profile_parser_agent_auto",
                entity_type="candidate",
                entity_id=candidate_id,
                details=json.dumps({
                    "parser": parsed_data.get("_parser", "unknown"), 
                    "auto_triggered": True,
                    "job_id": job_id
                }),
            )
            db.add(audit)
            await db.commit()

            # Publish event with complete data for Stage 3
            await event_bus.publish(
                EventTopics.PROFILE_PARSED,
                {
                    "candidate_id": candidate_id,
                    "job_id": job_id,
                    "application_id": application.id if application else None,
                    "name": candidate.name,
                    "email": candidate.email,
                    "phone": candidate.phone,
                    "location": candidate.location,
                    "skills": parsed_data.get("skills", []),
                    "experience_years": parsed_data.get("experience_years", 0),
                    "education": parsed_data.get("education", []),
                    "work_history": parsed_data.get("work_history", []),
                    "resume_path": candidate.resume_url,
                    "parser": parsed_data.get("_parser"),
                    "current_role": candidate.current_role
                },
                agent="profile_parser_agent_auto",
            )

            print(f"✅ Auto-parsed resume for {candidate.name} using {parsed_data.get('_parser', 'unknown')}")

        except Exception as e:
            await db.rollback()
            print(f"❌ Auto-parsing error for candidate {candidate_id}: {e}")


@router.post("/parse-resume", response_model=dict)
async def parse_resume(
    req: ParseRequest,
    db: AsyncSession = Depends(get_db),
    user: dict = Depends(get_current_user),
):
    """Parse a candidate's resume using LlamaIndex + Groq."""
    result = await db.execute(select(Candidate).where(Candidate.id == req.candidate_id))
    candidate = result.scalar_one_or_none()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")

    # Get application if exists
    app_result = await db.execute(select(Application).where(Application.candidate_id == req.candidate_id))
    application = app_result.scalar_one_or_none()
    job_id = application.job_id if application else candidate.job_id

    if not candidate.resume_url or not os.path.exists(candidate.resume_url):
        raise HTTPException(status_code=400, detail="No resume file found for this candidate")

    # Step 1: Extract text
    raw_text = extract_text(candidate.resume_url)
    if not raw_text:
        raise HTTPException(status_code=400, detail="Could not extract text from resume")

    # Step 2: Parse with LlamaIndex
    parsed_data = await _parse_with_llamaindex(raw_text)

    # Step 3: Update candidate record
    candidate.raw_resume_text = raw_text
    candidate.parsed_data = json.dumps(parsed_data)
    candidate.name = parsed_data.get("name") or candidate.name
    candidate.email = parsed_data.get("email") or candidate.email
    candidate.phone = parsed_data.get("phone") or candidate.phone
    candidate.location = parsed_data.get("location") or candidate.location
    candidate.current_role = parsed_data.get("current_role") or candidate.current_role
    candidate.experience_years = parsed_data.get("experience_years") or candidate.experience_years
    candidate.skills = json.dumps(parsed_data.get("skills", []))
    candidate.education = json.dumps(parsed_data.get("education", []))
    candidate.work_history = json.dumps(parsed_data.get("work_history", []))
    candidate.status = "parsed"
    candidate.updated_at = datetime.utcnow()
    
    # CRITICAL FIX: Set job_id on candidate for Stage 3 screening
    if job_id:
        candidate.job_id = job_id
        print(f"✅ Set job_id {job_id} on candidate {req.candidate_id}")
    else:
        print(f"⚠️ No job_id found for candidate {req.candidate_id} - Stage 3 screening will fail")

    # Audit log
    audit = AuditLog(
        id=generate_id(),
        event_type=EventTopics.PROFILE_PARSED,
        agent_name="profile_parser_agent",
        entity_type="candidate",
        entity_id=req.candidate_id,
        details=json.dumps({"parser": parsed_data.get("_parser", "unknown"), "job_id": job_id}),
    )
    db.add(audit)
    await db.commit()

    # Publish event with complete data for Stage 3
    await event_bus.publish(
        EventTopics.PROFILE_PARSED,
        {
            "candidate_id": req.candidate_id,
            "job_id": job_id,
            "application_id": application.id if application else None,
            "name": candidate.name,
            "email": candidate.email,
            "phone": candidate.phone,
            "location": candidate.location,
            "skills": parsed_data.get("skills", []),
            "experience_years": parsed_data.get("experience_years", 0),
            "education": parsed_data.get("education", []),
            "work_history": parsed_data.get("work_history", []),
            "resume_path": candidate.resume_url,
            "parser": parsed_data.get("_parser"),
            "current_role": candidate.current_role
        },
        agent="profile_parser_agent",
    )

    return {
        "success": True,
        "candidate_id": req.candidate_id,
        "job_id": job_id,
        "parsed_data": parsed_data,
        "parser_used": parsed_data.get("_parser", "unknown"),
        "message": f"Resume parsed successfully using {parsed_data.get('_parser', 'unknown')}",
    }


@router.post("/parse-batch", response_model=dict)
async def parse_batch(
    req: ParseBatchRequest,
    db: AsyncSession = Depends(get_db),
    user: dict = Depends(get_current_user),
):
    """Parse multiple resumes in batch."""
    if req.parse_all_unparsed:
        result = await db.execute(
            select(Candidate).where(Candidate.status.in_(["uploaded", "new"]))
        )
        candidates = result.scalars().all()
        candidate_ids = [c.id for c in candidates]
    else:
        candidate_ids = req.candidate_ids

    results = []
    for cid in candidate_ids:
        try:
            parse_req = ParseRequest(candidate_id=cid)
            result = await parse_resume(parse_req, db, user)
            results.append({"candidate_id": cid, "status": "success", "parser": result.get("parser_used")})
        except Exception as e:
            results.append({"candidate_id": cid, "status": "failed", "error": str(e)})

    return {
        "success": True,
        "total": len(candidate_ids),
        "parsed": sum(1 for r in results if r["status"] == "success"),
        "failed": sum(1 for r in results if r["status"] == "failed"),
        "results": results,
    }
